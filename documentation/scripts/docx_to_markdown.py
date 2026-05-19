#!/usr/bin/env python3
"""
Convert «Пояснительная записка.docx» to Markdown sections with extracted images.

Usage (from repository root):
    python documentation/scripts/docx_to_markdown.py
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import unicodedata
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator, Optional, Union

import yaml
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = REPO_ROOT / "documentation" / "Пояснительная записка.docx"
OUTPUT_DIR = REPO_ROOT / "documentation" / "пояснительная-записка"
ASSETS_DIR = OUTPUT_DIR / "assets"

FIGURE_RE = re.compile(
    r"Рисунок\s+(\d+|[Хх]{2,}|\d+)\s*[–—\-]\s*(.+)",
    re.IGNORECASE,
)
DRAFT_FIGURE_RE = re.compile(r"рисунок\s+х{2,}", re.IGNORECASE)

TRANSLIT = {
    "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "yo",
    "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
    "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
    "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "sch",
    "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
}

STRUCTURAL_TITLES = {
    "РЕФЕРАТ": ("01-referat", "Реферат"),
    "СОДЕРЖАНИЕ": ("02-soderzhanie", "Содержание"),
    "ВВЕДЕНИЕ": ("03-vvedenie", "Введение"),
    "ЗАКЛЮЧЕНИЕ": ("90-zaklyuchenie", "Заключение"),
}

STRUCTURAL_PREFIXES = [
    ("СПИСОК ИСПОЛЬЗОВАННЫХ", "91-istochniki", "Список использованных источников"),
    ("ПРИЛОЖЕНИЕ А", "92-prilozhenie-a", "Приложение А"),
    ("ПРИЛОЖЕНИЕ Б", "93-prilozhenie-b", "Приложение Б"),
]

MAJOR_SECTIONS = [
    (241, "04-predmetnaya-oblast", "Описание и анализ предметной области"),
    (428, "05-proektirovanie", "Проектирование системы"),
    (723, "06-realizaciya", "Реализация системы"),
]


@dataclass
class FigureRecord:
    id: str
    number: Optional[int]
    caption: str
    full_caption: str
    file: str
    section: str
    keywords: list[str]
    context_snippet: str
    draft: bool = False


@dataclass
class Section:
    section_id: str
    title: str
    lines: list[str] = field(default_factory=list)
    figure_count: int = 0

    @property
    def filename(self) -> str:
        return f"{self.section_id}.md"


@dataclass
class ImageAsset:
    r_id: str
    fig_id: str
    filename: str
    rel_path: str  # assets/fig-...


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\t", " ")
    text = re.sub(r" +", " ", text)
    return text.strip()


def slugify(text: str, max_len: int = 50) -> str:
    text = normalize_text(text).lower()
    result: list[str] = []
    for ch in text:
        low = ch.lower()
        if low in TRANSLIT:
            result.append(TRANSLIT[low])
        elif ch.isascii() and ch.isalnum():
            result.append(ch)
        elif ch in " -_":
            result.append("-")
    slug = "".join(result)
    slug = re.sub(r"-+", "-", slug).strip("-")
    if not slug:
        slug = "image"
    return slug[:max_len].rstrip("-")


def parse_figure_caption(text: str) -> Optional[tuple[Optional[int], str, str]]:
    text = normalize_text(text)
    m = FIGURE_RE.search(text)
    if not m:
        return None
    num_raw = m.group(1)
    desc = normalize_text(m.group(2))
    if re.fullmatch(r"[хХ]+", num_raw):
        number = None
    else:
        try:
            number = int(num_raw)
        except ValueError:
            number = None
    full = f"Рисунок {num_raw} – {desc}"
    return number, desc, full


def iter_block_items(parent: Union[DocumentObject, _Cell]) -> Iterator[Union[Paragraph, Table]]:
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//a:blip"))


def get_embed_ids(paragraph: Paragraph) -> list[str]:
    ids: list[str] = []
    for blip in paragraph._element.xpath(".//a:blip"):
        embed = blip.get(qn("r:embed"))
        if embed:
            ids.append(embed)
    return ids


def runs_to_markdown(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text.replace("\xa0", " ")
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    joined = "".join(parts)
    return re.sub(r" +", " ", joined).strip()


def paragraph_style_heading_level(style_name: str) -> Optional[int]:
    sn = style_name or ""
    if sn == "МР_Структурный элемент":
        return 1
    if sn == "МР_Раздел":
        return 2
    if sn == "МР_Подраздел":
        return 3
    if sn == "МР_Параграф":
        return 4
    if "Heading 1" in sn:
        return 1
    if "Heading 2" in sn:
        return 2
    if "заголовок" in sn.lower():
        if "1" in sn:
            return 1
        if "2" in sn:
            return 2
        if "3" in sn:
            return 3
    return None


def table_to_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [normalize_text(cell.text) for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")
    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        escaped = [c.replace("|", "\\|") for c in row]
        lines.append("| " + " | ".join(escaped) + " |")
    return "\n".join(lines)


def convert_wmf_to_png(wmf_path: Path, png_path: Path) -> bool:
    for cmd_name, args in [
        ("magick", ["convert", str(wmf_path), str(png_path)]),
        ("convert", [str(wmf_path), str(png_path)]),
        ("inkscape", [str(wmf_path), "--export-type=png", f"--export-filename={png_path}"]),
    ]:
        exe = shutil.which(cmd_name)
        if not exe:
            continue
        try:
            cmd = [exe, *args] if cmd_name != "magick" else [exe, "convert", str(wmf_path), str(png_path)]
            subprocess.run(cmd, check=True, capture_output=True, timeout=60)
            if png_path.exists():
                return True
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return False


def save_image_blob(
    blob: bytes,
    ext: str,
    dest_stem: Path,
) -> tuple[str, Path]:
    ext = ext.lower().lstrip(".")
    if ext in ("jpg", "jpeg"):
        out = dest_stem.with_suffix(".jpg")
        out.write_bytes(blob)
        return out.name, out
    if ext == "wmf":
        wmf_path = dest_stem.with_suffix(".wmf")
        wmf_path.write_bytes(blob)
        png_path = dest_stem.with_suffix(".png")
        if convert_wmf_to_png(wmf_path, png_path):
            wmf_path.unlink(missing_ok=True)
            return png_path.name, png_path
        try:
            from PIL import Image

            with Image.open(wmf_path) as im:
                im.save(png_path, "PNG")
            wmf_path.unlink(missing_ok=True)
            return png_path.name, png_path
        except Exception:
            return wmf_path.name, wmf_path
    out = dest_stem.with_suffix(f".{ext or 'png'}")
    out.write_bytes(blob)
    return out.name, out


def find_caption_in_window(texts: list[str], index: int) -> Optional[tuple[int, Optional[int], str, str]]:
    for offset in (0, 1, 2, 3, -1, -2, -3):
        j = index + offset
        if 0 <= j < len(texts):
            parsed = parse_figure_caption(texts[j])
            if parsed:
                return j, parsed[0], parsed[1], parsed[2]
    return None


def detect_section_start(paragraph: Paragraph, para_index: int) -> Optional[tuple[str, str]]:
    text = normalize_text(paragraph.text)
    style = paragraph.style.name if paragraph.style else ""

    if para_index < 222:
        if para_index == 0:
            return "00-zadanie", "Задание на курсовой проект"
        return None

    upper = text.upper().split("\n")[0].strip()
    for key, (sid, title) in STRUCTURAL_TITLES.items():
        if upper == key or upper.startswith(key):
            return sid, title

    for prefix, sid, title in STRUCTURAL_PREFIXES:
        if upper.startswith(prefix):
            return sid, title

    if style == "МР_Раздел" and text and len(text) < 100:
        for start_idx, sid, title in MAJOR_SECTIONS:
            if para_index >= start_idx and (text == title or text.startswith(title[:20])):
                return sid, title

    return None


def resolve_section_id(para_index: int, explicit: Optional[tuple[str, str]]) -> tuple[str, str]:
    if explicit:
        return explicit
    if para_index < 222:
        return "00-zadanie", "Задание на курсовой проект"
    if para_index < 233:
        return "01-referat", "Реферат"
    if para_index < 241:
        return "03-vvedenie", "Введение"
    if para_index < 428:
        return "04-predmetnaya-oblast", "Описание и анализ предметной области"
    if para_index < 723:
        return "05-proektirovanie", "Проектирование системы"
    if para_index < 810:
        return "06-realizaciya", "Реализация системы"
    if para_index < 817:
        return "90-zaklyuchenie", "Заключение"
    if para_index < 852:
        return "91-istochniki", "Список использованных источников"
    if para_index < 876:
        return "92-prilozhenie-a", "Приложение А"
    return "93-prilozhenie-b", "Приложение Б"


class DocxConverter:
    def __init__(self, docx_path: Path, output_dir: Path) -> None:
        self.docx_path = docx_path
        self.output_dir = output_dir
        self.assets_dir = output_dir / "assets"
        self.doc = Document(str(docx_path))
        self.sections: dict[str, Section] = {}
        self.figures: list[FigureRecord] = []
        self.rid_to_asset: dict[str, ImageAsset] = {}
        self.caption_para_indices: set[int] = set()
        self.figure_seq = 0
        self.para_texts: list[str] = []
        self.para_blocks: list[Paragraph] = []
        self.block_meta: list[tuple[str, object, int]] = []  # kind, block, para_index

    def _get_section(self, section_id: str, title: str) -> Section:
        if section_id not in self.sections:
            self.sections[section_id] = Section(section_id=section_id, title=title)
        return self.sections[section_id]

    def _collect_blocks(self) -> None:
        para_idx = 0
        for block in iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                self.para_blocks.append(block)
                self.para_texts.append(normalize_text(block.text))
                self.block_meta.append(("paragraph", block, para_idx))
                para_idx += 1
            else:
                self.block_meta.append(("table", block, -1))

    def _extract_and_register_image(
        self,
        paragraph: Paragraph,
        para_index: int,
        embed_id: str,
        section_id: str,
    ) -> ImageAsset:
        if embed_id in self.rid_to_asset:
            return self.rid_to_asset[embed_id]

        part = paragraph.part.related_parts.get(embed_id)
        if part is None:
            raise RuntimeError(f"Missing image part for rId={embed_id}")

        blob = part.blob
        ext = part.content_type.split("/")[-1]
        if ext == "x-wmf":
            ext = "wmf"

        cap_result = find_caption_in_window(self.para_texts, para_index)
        if cap_result:
            cap_idx, num, desc, full_caption = cap_result
            self.caption_para_indices.add(cap_idx)
        else:
            ctx_parts = []
            for j in range(max(0, para_index - 2), min(len(self.para_texts), para_index + 3)):
                if self.para_texts[j] and not parse_figure_caption(self.para_texts[j]):
                    ctx_parts.append(self.para_texts[j])
            desc = ctx_parts[0][:80] if ctx_parts else f"image-{self.figure_seq + 1}"
            num = None
            full_caption = f"Рисунок – {desc}"
            cap_idx = -1

        self.figure_seq += 1
        fig_num_str = f"{self.figure_seq:02d}"
        slug = slugify(desc)
        fig_id = f"fig-{fig_num_str}"
        stem = self.assets_dir / f"{fig_id}-{slug}"

        filename, _ = save_image_blob(blob, ext, stem)
        rel_path = f"assets/{filename}"
        asset = ImageAsset(r_id=embed_id, fig_id=fig_id, filename=filename, rel_path=rel_path)
        self.rid_to_asset[embed_id] = asset

        draft = DRAFT_FIGURE_RE.search(full_caption) is not None
        keywords = list({w.lower() for w in re.findall(r"[а-яёa-z0-9]{3,}", desc, re.I)} | {slug.replace("-", " ")})
        keywords = [k for k in keywords if k][:12]

        snippet = ""
        for j in range(max(0, para_index - 2), min(len(self.para_texts), para_index + 2)):
            t = self.para_texts[j]
            if t and j != cap_idx and not parse_figure_caption(t):
                snippet = t[:200]
                break

        self.figures.append(
            FigureRecord(
                id=fig_id,
                number=num,
                caption=desc,
                full_caption=full_caption,
                file=rel_path,
                section=section_id,
                keywords=sorted(keywords),
                context_snippet=snippet,
                draft=draft,
            )
        )
        return asset

    def _render_image_md(self, asset: ImageAsset, full_caption: str) -> str:
        alt = full_caption
        lines = [
            f"<!-- fig-id: {asset.fig_id} -->",
            f"![{alt}](./{asset.rel_path})",
            "",
            f"*{full_caption}*",
            "",
        ]
        return "\n".join(lines)

    def convert(self) -> None:
        self.assets_dir.mkdir(parents=True, exist_ok=True)
        for old in self.assets_dir.glob("*"):
            if old.is_file():
                old.unlink()

        self._collect_blocks()

        current_section_id = "00-zadanie"
        current_title = "Задание на курсовой проект"

        for kind, block, para_index in self.block_meta:
            if kind == "table":
                section = self._get_section(current_section_id, current_title)
                md = table_to_markdown(block)
                if md:
                    section.lines.append(md)
                    section.lines.append("")
                continue

            paragraph: Paragraph = block
            text = self.para_texts[para_index]

            new_sec = detect_section_start(paragraph, para_index)
            if new_sec:
                current_section_id, current_title = new_sec
            elif para_index >= 222:
                current_section_id, current_title = resolve_section_id(para_index, None)

            section = self._get_section(current_section_id, current_title)

            if para_index in self.caption_para_indices:
                continue

            embed_ids = get_embed_ids(paragraph)
            if embed_ids:
                cap_result = find_caption_in_window(self.para_texts, para_index)
                if cap_result:
                    _, _, _, full_caption = cap_result
                else:
                    full_caption = f"Рисунок – {text[:80]}" if text else "Рисунок"
                for eid in embed_ids:
                    asset = self._extract_and_register_image(
                        paragraph, para_index, eid, current_section_id
                    )
                    section.lines.append(self._render_image_md(asset, full_caption))
                    section.figure_count += 1
                # if paragraph is only caption text, skip rest
                if text and parse_figure_caption(text):
                    continue

            if not text:
                if section.lines and section.lines[-1] != "":
                    section.lines.append("")
                continue

            if parse_figure_caption(text):
                continue

            heading_level = paragraph_style_heading_level(
                paragraph.style.name if paragraph.style else ""
            )
            md_text = runs_to_markdown(paragraph) or text

            style_name = paragraph.style.name if paragraph.style else ""
            if "Список маркированный" in style_name or style_name == "МР_Список маркированный":
                section.lines.append(f"- {md_text}")
            elif "список нумерованный" in style_name.lower() or "нумерованный" in style_name.lower():
                section.lines.append(f"1. {md_text}")
            elif heading_level:
                header = f"{'#' * (heading_level + 1)} {md_text}"
                title_norm = re.sub(r"\s+", "", section.title.lower())
                text_norm = re.sub(r"\s+", "", md_text.lower())
                if text_norm != title_norm:
                    section.lines.append(header)
            else:
                section.lines.append(md_text)
            section.lines.append("")

        self._write_sections()
        self._write_manifest()
        self._write_index()

    def _write_sections(self) -> None:
        for section in sorted(self.sections.values(), key=lambda s: s.section_id):
            path = self.output_dir / section.filename
            front_matter = (
                "---\n"
                f"title: \"{section.title}\"\n"
                f"section_id: {section.section_id}\n"
                "source: docx\n"
                "---\n\n"
            )
            body = "\n".join(section.lines)
            body = re.sub(r"\n{3,}", "\n\n", body).strip() + "\n"
            path.write_text(front_matter + f"# {section.title}\n\n" + body, encoding="utf-8")

    def _write_manifest(self) -> None:
        data = [
            {
                "id": f.id,
                "number": f.number,
                "caption": f.caption,
                "full_caption": f.full_caption,
                "file": f.file,
                "section": f.section,
                "keywords": f.keywords,
                "context_snippet": f.context_snippet,
                "draft": f.draft,
            }
            for f in self.figures
        ]
        manifest_path = self.output_dir / "manifest.yaml"
        with manifest_path.open("w", encoding="utf-8") as fh:
            yaml.dump(
                {"figures": data, "generated_at": datetime.now(timezone.utc).isoformat()},
                fh,
                allow_unicode=True,
                sort_keys=False,
                default_flow_style=False,
            )

    def _write_index(self) -> None:
        lines = [
            "# Пояснительная записка (Markdown)",
            "",
            f"Сконвертировано из `{self.docx_path.name}` "
            f"({datetime.now(timezone.utc).strftime('%Y-%m-%d')}).",
            "",
            "Оригинал: [Пояснительная записка.docx](../Пояснительная%20записка.docx)",
            "",
            "## Разделы",
            "",
        ]
        for section in sorted(self.sections.values(), key=lambda s: s.section_id):
            lines.append(f"- [{section.title}](./{section.filename})")

        lines.extend(["", "## Рисунки", "", "| № | Подпись | Файл | Раздел |", "|---|---------|------|--------|"])
        for f in self.figures:
            num = str(f.number) if f.number is not None else "—"
            draft = " (draft)" if f.draft else ""
            lines.append(
                f"| {num} | {f.caption}{draft} | [{f.file}](./{f.file}) | {f.section} |"
            )
        lines.append("")
        (self.output_dir / "index.md").write_text("\n".join(lines), encoding="utf-8")


def validate_output(output_dir: Path, doc: Document) -> list[str]:
    issues: list[str] = []
    assets = list((output_dir / "assets").glob("*"))
    if len(assets) < 40:
        issues.append(f"Expected ~43 assets, found {len(assets)}")

    md_files = list(output_dir.glob("*.md"))
    md_files = [f for f in md_files if f.name != "index.md" and f.name != "README.md"]
    if len(md_files) < 5:
        issues.append(f"Too few section files: {len(md_files)}")

    doc_words = sum(len(p.text.split()) for p in doc.paragraphs)
    md_words = 0
    for mf in md_files:
        md_words += len(mf.read_text(encoding="utf-8").split())
    ratio = md_words / doc_words if doc_words else 0
    if ratio < 0.85 or ratio > 1.15:
        issues.append(f"Word count ratio {ratio:.2f} (doc={doc_words}, md={md_words})")

    for mf in output_dir.rglob("*.md"):
        text = mf.read_text(encoding="utf-8")
        for m in re.finditer(r"!\[[^\]]*\]\((\./[^)]+)\)", text):
            rel = m.group(1).lstrip("./")
            if not (output_dir / rel).exists():
                issues.append(f"Broken image link in {mf.name}: {rel}")

    return issues


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert DOCX explanatory note to Markdown.")
    parser.add_argument("--docx", type=Path, default=DOCX_PATH)
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR)
    parser.add_argument("--validate", action="store_true", default=True)
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"DOCX not found: {args.docx}", file=sys.stderr)
        return 1

    args.output.mkdir(parents=True, exist_ok=True)
    converter = DocxConverter(args.docx, args.output)
    converter.convert()
    print(f"Wrote {len(converter.sections)} sections, {len(converter.figures)} figures to {args.output}")

    if args.validate:
        doc = Document(str(args.docx))
        issues = validate_output(args.output, doc)
        if issues:
            print("Validation warnings:")
            for i in issues:
                print(f"  - {i}")
        else:
            print("Validation OK")

    return 0


if __name__ == "__main__":
    sys.exit(main())
